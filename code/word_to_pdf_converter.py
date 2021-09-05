import sys
from docx2pdf import convert
from docx import Document
import re 
import os
import subprocess


def job_title_cleaning(job_title):
    job_title = job_title.title()

    if job_title[0].lower() in ['u','e','o','a','i']:
        job_title = 'an ' + job_title
    else: 
        job_title = 'a ' + job_title
    
    return job_title

# ref: https://www.quora.com/How-can-I-find-and-replace-text-in-a-Word-document-using-Python
def docx_replace_regex(doc_obj, regex , replace): 
    company_team_flag = False

    for p in doc_obj.paragraphs: 
        if regex.search(p.text): 
            inline = p.runs 
            # Loop added to work with runs (strings with same style) 
            for i in range(len(inline)): 
                if regex.search(inline[i].text): 
                    text = regex.sub(replace, inline[i].text) 
                    inline[i].text = text
                                        
                    if text == 'COMPANYTEAM':
                        company_team_flag = True
                        break
            if company_team_flag == True:
                break
 
    for table in doc_obj.tables: 
        for row in table.rows: 
            for cell in row.cells: 
                docx_replace_regex(cell, regex , replace)

def replace_company_name(company_team, company_name, job_title):
    regex1 = re.compile(r"COMPANYTEAM") 
    replace1 = company_team

    regex2 = re.compile(r"COMPANYNAME") 
    replace2 = company_name

    regex3 = re.compile(r"a JOBPOSITION")    
    replace3 = job_title
    
    filename = "./CoverLetter-Max.docx"

    doc = Document(filename) 
    docx_replace_regex(doc, regex1 , replace1)
    docx_replace_regex(doc, regex2 , replace2) 
    docx_replace_regex(doc, regex3 , replace3) 
    doc.save('./CoverLetter-Max.docx') 

def replace_company_name_original(company_team, company_name, job_title):
    regex1 = re.compile(company_team) 
    replace1 = "COMPANYTEAM"

    regex2 = re.compile(company_name) 
    replace2 = "COMPANYNAME"
    
    regex3 = re.compile(job_title) 
    replace3 = "a JOBPOSITION"
    
    filename = "./CoverLetter-Max.docx"

    doc = Document(filename) 
    docx_replace_regex(doc, regex1 , replace1) 
    docx_replace_regex(doc, regex2 , replace2) 
    docx_replace_regex(doc, regex3 , replace3) 

    doc.save('./CoverLetter-Max.docx') 
    pass

def doc_to_pdf_converter():
    # Current directory
    doc_path = os.getcwd()+'\CoverLetter-Max.docx'
    # Parent directory
    pdf_path = os.path.abspath(os.getcwd() + "/../") + '\CoverLetter-Max.pdf'
    
    # print(doc_path)
    # print(pdf_path)
    convert(doc_path, pdf_path)

def pdf_open_check():
    pdf_path = os.path.abspath(os.getcwd() + "/../") + '\CoverLetter-Max.pdf'
    subprocess.Popen([pdf_path],shell=True)

def main():
    company_team = input("Enter your company's team: ")
    company_name = input("Enter your company's name: ")
    job_title = input("Enter your job title: ")
    job_title = job_title_cleaning(job_title)
    print(company_team, company_name, job_title)

    replace_company_name(company_team, company_name, job_title)
    doc_to_pdf_converter()
    replace_company_name_original(company_team, company_name, job_title)
    pdf_open_check()


if __name__ == "__main__":
    main()